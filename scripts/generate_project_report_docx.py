#!/usr/bin/env python3
"""Generate a clean, descriptive DOCX report with metrics, tables, and graphs."""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report_assets"
METRICS_ENSEMBLE = ROOT / "website" / "models" / "vit_ensemble_metrics.json"
METRICS_SINGLE = ROOT / "website" / "models" / "metrics.json"
CONFUSION = ROOT / "website" / "models" / "vit_ensemble_confusion.csv"
OUT_DOCX = DOCS / "Skin_Disease_Project_Report.docx"


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.size = Pt(15)
    elif level == 2:
        h.runs[0].font.size = Pt(12)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h

    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            cells[i].text = str(val)


def add_image_with_caption(doc: Document, image_path: Path, caption: str, width_inches: float = 6.2) -> None:
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width_inches))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
    else:
        add_para(doc, f"[Missing figure: {image_path.name}]", italic=True)


def confusion_rows(path: Path) -> tuple[list[str], list[list[str]]]:
    with path.open(encoding="utf-8") as f:
        reader = csv.reader(f)
        header = next(reader)
        labels = header
        rows = [r for r in reader]
    return labels, rows


def main() -> int:
    metrics = json.loads(METRICS_ENSEMBLE.read_text(encoding="utf-8"))
    single = json.loads(METRICS_SINGLE.read_text(encoding="utf-8"))

    doc = Document()
    set_default_style(doc)

    # Cover page
    add_para(doc, "Skin Disease Classifier", bold=True, center=True)
    add_para(doc, "Student Group Project Report", bold=True, center=True)
    add_para(doc, "Applied Image Processing", italic=True, center=True)
    add_para(doc, "Academic Year: 2025-2026", center=True)
    add_para(doc, "Submission Date: 21 April 2026", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Prepared by: Group 19", center=True)
    add_para(doc, "Anurag Thakur 2023IMG-008", center=True)
    add_para(doc, "Baviskar Pratik Subhash 2023IMG-014", center=True)
    add_para(doc, "Kshitij Dhamanikar 2023IMG-029", center=True)
    add_para(doc, "Kushagra Agarwal 2023IMG-030", center=True)
    add_para(doc, "Vibhor Kumar 2023IMG-054", center=True)
    doc.add_page_break()

    # Abstract
    add_heading(doc, "Abstract", level=1)
    add_para(
        doc,
        "This report presents a complete machine learning workflow for automated skin disease image classification. "
        "The final system uses a Vision Transformer (ViT)-based weighted ensemble for 8-class prediction. "
        "Compared to a single promoted ViT baseline, the final ensemble improves overall accuracy and macro-level metrics substantially. "
        "The deployment-ready artifacts were integrated into a Flask web application runtime and validated through class-wise metrics, "
        "confusion analysis, and leakage-aware evaluation checks.",
    )
    add_para(
        doc,
        f"Final ensemble metrics on the main test split: Accuracy={metrics['accuracy']:.4f}, "
        f"Macro Precision={metrics['macro_precision']:.4f}, Macro Recall={metrics['macro_recall']:.4f}.",
        bold=True,
    )

    add_heading(doc, "Keywords", level=2)
    add_para(doc, "Skin disease classification, Vision Transformer, Ensemble learning, Medical image analysis, Deep learning")

    # Introduction
    add_heading(doc, "1. Introduction", level=1)
    add_para(
        doc,
        "Skin disorders are common and often require timely assessment. Automated classification tools can support screening, "
        "education, and decision support when designed with clear limitations. This project focuses on building a practical, "
        "reproducible image classification pipeline under CPU-only constraints while maintaining strong model quality.",
    )

    add_heading(doc, "1.1 Problem Statement", level=2)
    add_para(
        doc,
        "Given an input skin image, classify the disease category among 8 predefined classes: acne, bcc, bkl, eczema, ak, melanoma, psoriasis, and tinea.",
    )

    add_heading(doc, "1.2 Objectives", level=2)
    add_bullets(
        doc,
        [
            "Build CPU-friendly ViT training and evaluation workflows.",
            "Improve model quality through controlled ensemble optimization.",
            "Integrate final model configuration into application runtime.",
            "Produce clear metrics, plots, and report artifacts for academic submission.",
            "Validate performance with leakage-aware checks for reliability.",
        ],
    )

    add_heading(doc, "1.3 Scope and Limitations", level=2)
    add_para(
        doc,
        "The system is intended for educational and research use and is not a replacement for clinical diagnosis. "
        "Model performance is dataset-dependent and can degrade on unseen data domains.",
    )

    # Data and setup
    add_heading(doc, "2. Dataset and Experimental Setup", level=1)
    add_para(
        doc,
        "Experiments were run on curated split folders with balanced class representation in the test split. "
        "The final evaluated split consists of 184 total images with 23 images per class.",
    )

    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Primary evaluation split", "datasets/hf_hybrid_150/splits/test"],
            ["Number of classes", "8"],
            ["Total test samples", str(metrics["samples"])],
            ["Samples per class", "23"],
            ["Runtime environment", "CPU-only training/evaluation"],
            ["Model family", "Vision Transformer (ViT)"],
        ],
    )

    add_heading(doc, "2.1 Leakage-Aware Evaluation Note", level=2)
    add_para(
        doc,
        "Because some dataset versions may share source images, leakage-safe subset checks were used to validate whether gains persist "
        "after excluding overlapping train/validation references from expanded dataset variants.",
    )

    add_heading(doc, "2.2 Experimental Protocol", level=2)
    add_bullets(
        doc,
        [
            "Train multiple ViT checkpoints with varied augmentation and regularization profiles.",
            "Evaluate checkpoints on common split for comparability.",
            "Search ensemble weights for improved accuracy and stability.",
            "Promote best-performing ensemble into runtime config.",
        ],
    )

    # Architecture
    add_heading(doc, "3. Model Architecture and Pipeline", level=1)
    add_para(
        doc,
        "The final predictor is a weighted softmax ensemble over four ViT checkpoints. "
        "Each member contributes probabilities according to fixed weights, and the class with maximum weighted probability is selected.",
    )

    add_table(
        doc,
        ["Ensemble Member", "Model Path", "Weight"],
        [
            ["strict80_tiny160", "artifacts/vit_strict80_tiny160/best_model.pt", "0.50"],
            ["220_tiny224_lowaug", "artifacts/vit_220_tiny224_lowaug/best_model.pt", "0.40"],
            ["150_tiny224_lowaug", "artifacts/vit_150_tiny224_lowaug/best_model.pt", "0.05"],
            ["150_tiny160", "artifacts/vit_150_tiny160/best_model.pt", "0.05"],
        ],
    )

    add_image_with_caption(doc, ASSETS / "ensemble_weights_pie.png", "Figure 1: Final ensemble member weight distribution.")

    add_heading(doc, "3.1 Deployment Path", level=2)
    add_bullets(
        doc,
        [
            "Runtime config: website/models/vit_ensemble.json",
            "Metrics export: website/models/vit_ensemble_metrics.json",
            "Confusion export: website/models/vit_ensemble_confusion.csv",
            "Serving interface: Flask application inference route",
        ],
    )

    # Results
    add_heading(doc, "4. Final Quantitative Results", level=1)
    add_image_with_caption(doc, ASSETS / "kpi_summary.png", "Figure 2: KPI summary of final ensemble.")

    add_table(
        doc,
        ["Metric", "Final Ensemble", "Single ViT Baseline", "Absolute Gain"],
        [
            ["Accuracy", f"{metrics['accuracy']:.4f}", f"{single['accuracy']:.4f}", f"+{metrics['accuracy']-single['accuracy']:.4f}"],
            [
                "Macro Precision",
                f"{metrics['macro_precision']:.4f}",
                f"{single['macro_precision']:.4f}",
                f"+{metrics['macro_precision']-single['macro_precision']:.4f}",
            ],
            [
                "Macro Recall",
                f"{metrics['macro_recall']:.4f}",
                f"{single['macro_recall']:.4f}",
                f"+{metrics['macro_recall']-single['macro_recall']:.4f}",
            ],
        ],
    )

    add_para(doc, "The ensemble gives a clear and consistent improvement over the single baseline model across all primary aggregate metrics.")

    add_heading(doc, "4.1 Class-wise Metrics", level=2)
    per_rows = []
    for cls in metrics["class_order"]:
        entry = metrics["per_class"][cls]
        per_rows.append([cls, f"{entry['precision']:.4f}", f"{entry['recall']:.4f}", str(entry["support"])])

    add_table(doc, ["Class", "Precision", "Recall", "Support"], per_rows)

    # Analysis
    add_heading(doc, "5. Diagnostic Analysis", level=1)
    add_image_with_caption(doc, ASSETS / "precision_recall_by_class.png", "Figure 3: Precision and recall by class.")
    add_image_with_caption(doc, ASSETS / "f1_by_class.png", "Figure 4: Per-class F1 score.")

    add_heading(doc, "5.1 Key Findings", level=2)
    add_bullets(
        doc,
        [
            "Acne and BCC show strong recall and stable precision.",
            "AK and Melanoma maintain balanced performance.",
            "Eczema remains the most difficult class in recall.",
            "Tinea has high precision but lower recall, suggesting conservative prediction behavior.",
        ],
    )

    # Confusion and distribution
    add_heading(doc, "6. Confusion and Prediction Distribution", level=1)
    add_image_with_caption(doc, ASSETS / "confusion_matrix_heatmap.png", "Figure 5: Confusion matrix heatmap.")
    add_image_with_caption(doc, ASSETS / "true_vs_predicted_counts.png", "Figure 6: True vs predicted class counts.")

    add_heading(doc, "6.1 Confusion Patterns", level=2)
    add_bullets(
        doc,
        [
            "Eczema has notable confusion with Psoriasis and Acne.",
            "AK occasionally overlaps with BCC.",
            "Tinea has occasional confusion with Psoriasis and Acne.",
            "Acne reached perfect recall in the final run.",
        ],
    )

    add_heading(doc, "6.2 Raw Confusion Table", level=2)
    conf_header, conf_rows = confusion_rows(CONFUSION)
    add_table(doc, conf_header, conf_rows)

    # Risks and conclusion
    add_heading(doc, "7. Risks, Ethics, and Limitations", level=1)
    add_bullets(
        doc,
        [
            "Performance remains below 90% target under CPU constraints.",
            "Visually similar classes can still produce boundary errors.",
            "Dataset overlap management remains critical for fair benchmarking.",
            "Uncertainty calibration and abstention are not yet implemented.",
        ],
    )

    add_heading(doc, "7.1 Ethical Note", level=2)
    add_para(
        doc,
        "This model is intended for educational and research use only. It is not a medical device and should not be used as a standalone diagnostic tool.",
    )

    add_heading(doc, "8. Conclusion and Future Work", level=1)
    add_para(
        doc,
        "The project successfully delivered an integrated ViT ensemble pipeline with strong improvements over baseline. "
        "The generated artifacts include deployable configuration files, quantitative reports, confusion diagnostics, and reproducible scripts.",
    )

    add_bullets(
        doc,
        [
            "Train higher-capacity ViT backbones on strict non-overlap datasets.",
            "Introduce confidence calibration for safer prediction handling.",
            "Perform external benchmark validation for robustness estimation.",
            "Add monitoring hooks for model drift in deployment.",
        ],
    )

    add_heading(doc, "Appendix A: Artifact References", level=2)
    add_table(
        doc,
        ["Artifact", "Purpose"],
        [
            ["website/models/vit_ensemble.json", "Final runtime ensemble configuration"],
            ["website/models/vit_ensemble_metrics.json", "Primary final metrics"],
            ["website/models/vit_ensemble_confusion.csv", "Confusion matrix values"],
            ["docs/report_assets/*.png", "Figures used in report and presentation"],
            ["scripts/generate_project_report_docx.py", "DOCX report generator"],
        ],
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"saved: {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
